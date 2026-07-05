"""
parser.py
---------
Handles extraction and cleaning of text from resumes (PDF/DOCX/TXT)
and plain-text job descriptions.
"""

import re
import io
from pathlib import Path
from typing import Union

import pdfplumber
import docx


def extract_text_from_pdf(file_path: Union[str, Path, io.BytesIO]) -> str:
    """Extract raw text from a PDF file (path or file-like object)."""
    text_chunks = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_path: Union[str, Path, io.BytesIO]) -> str:
    """Extract raw text from a DOCX file (path or file-like object)."""
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of tables, since many resumes use table layouts
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text(file_path: Union[str, Path], file_bytes: io.BytesIO = None) -> str:
    """
    Dispatch to the right extractor based on file extension.
    Accepts either a real path on disk, or a filename + in-memory bytes
    (useful when the file comes from a Streamlit file_uploader).
    """
    suffix = Path(file_path).suffix.lower()
    source = file_bytes if file_bytes is not None else file_path

    if suffix == ".pdf":
        return extract_text_from_pdf(source)
    elif suffix == ".docx":
        return extract_text_from_docx(source)
    elif suffix == ".txt":
        if file_bytes is not None:
            return file_bytes.read().decode("utf-8", errors="ignore")
        return Path(file_path).read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .docx, or .txt")


def clean_text(raw_text: str) -> str:
    """
    Normalize whitespace, strip bullet symbols/odd unicode, and
    collapse repeated blank lines so downstream NLP works on clean text.
    """
    text = raw_text.replace("\u2022", " ").replace("\uf0b7", " ")  # common bullet glyphs
    text = re.sub(r"[•▪●◦‣]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n", text)
    text = re.sub(r"[^\x00-\x7F]+", " ", text)  # drop stray non-ascii artifacts from PDF extraction
    return text.strip()
